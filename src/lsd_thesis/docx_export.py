from __future__ import annotations

import re
from pathlib import Path

from docx import Document as DocumentFactory
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

_IMAGE_RE = re.compile(r"^!\[(?P<alt>.*?)\]\((?P<body>.+)\)$")
_IMAGE_BODY_RE = re.compile(r"^(?P<path><[^>]+>|.+?)(?:\s+\"(?P<title>.*)\")?$")


def _clean_text(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text.strip()


def _configure_document(document: DocxDocument) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Times New Roman"


def _resolve_image_path(source_path: Path, image_path: str) -> Path:
    candidate = Path(image_path)
    if candidate.is_absolute():
        return candidate
    return (source_path.parent / candidate).resolve()


def _parse_image_target(body: str) -> tuple[str, str | None]:
    match = _IMAGE_BODY_RE.match(body.strip())
    if match is None:
        raise ValueError(f"Invalid markdown image target: {body!r}")
    path = match.group("path").strip()
    if path.startswith("<") and path.endswith(">"):
        path = path[1:-1].strip()
    return path, match.group("title")


def _add_heading(document: DocxDocument, text: str, level: int) -> None:
    if level <= 0:
        paragraph = document.add_paragraph(style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(text)
        return
    document.add_heading(text, level=min(level, 3))


def _add_paragraph(document: DocxDocument, text: str) -> None:
    if not text:
        return
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run(_clean_text(text))


def _add_bullet(document: DocxDocument, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run(_clean_text(text))


def _add_caption(document: DocxDocument, caption_text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(caption_text)


def _add_image(
    document: DocxDocument,
    source_path: Path,
    image_path: str,
    alt_text: str,
    caption_text: str | None = None,
) -> None:
    resolved_path = _resolve_image_path(source_path, image_path)
    if not resolved_path.exists():
        raise FileNotFoundError(f"Image not found: {resolved_path}")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(resolved_path), width=Inches(5.9))
    if caption_text:
        _add_caption(document, caption_text)


def markdown_to_docx(source_path: Path, output_path: Path) -> None:
    document = DocumentFactory()
    _configure_document(document)

    buffer: list[str] = []

    def flush_buffer() -> None:
        if buffer:
            _add_paragraph(document, " ".join(buffer))
            buffer.clear()

    for raw_line in source_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            flush_buffer()
            continue

        image_match = _IMAGE_RE.match(line)
        if image_match:
            flush_buffer()
            image_path, image_title = _parse_image_target(image_match.group("body"))
            _add_image(
                document,
                source_path,
                image_path,
                _clean_text(image_match.group("alt")),
                _clean_text(image_title) if image_title else None,
            )
            continue

        if line.startswith("#"):
            flush_buffer()
            heading_depth = len(line) - len(line.lstrip("#"))
            heading_text = _clean_text(line[heading_depth:].lstrip())
            _add_heading(document, heading_text, heading_depth - 1)
            continue

        if line.startswith("- "):
            flush_buffer()
            _add_bullet(document, line[2:])
            continue

        buffer.append(line)

    flush_buffer()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
