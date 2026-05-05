from __future__ import annotations

import base64
import os
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from lsd_thesis.docx_export import markdown_to_docx

_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7+X2QAAAAASUVORK5CYII="
)


def _non_empty_paragraph_texts(document: Document) -> list[str]:
    return [paragraph.text for paragraph in document.paragraphs if paragraph.text]


def _paragraph_with_text(document: Document, text: str):
    for paragraph in document.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise AssertionError(f"Paragraph with text {text!r} not found")


def test_markdown_to_docx_embeds_markdown_images_with_spaces_in_path(tmp_path: Path) -> None:
    source_dir = tmp_path / "notes"
    image_dir = source_dir / "figures"
    source_dir.mkdir()
    image_dir.mkdir()

    image_path = image_dir / "stage 1.png"
    image_path.write_bytes(_ONE_PIXEL_PNG)

    source_path = source_dir / "report.md"
    source_path.write_text(
        "# Thesis Report\n\n## Overview\n\n![Stage 1 figure](figures/stage 1.png)\n",
        encoding="utf-8",
    )

    output_path = tmp_path / "report.docx"
    markdown_to_docx(source_path, output_path)

    document = Document(output_path)
    assert output_path.exists()
    assert any(paragraph.style.name == "Title" and paragraph.text == "Thesis Report" for paragraph in document.paragraphs)
    assert any(paragraph.style.name == "Heading 1" and paragraph.text == "Overview" for paragraph in document.paragraphs)
    assert len(document.inline_shapes) == 1


def test_markdown_to_docx_raises_cleanly_for_missing_image_path(tmp_path: Path) -> None:
    source_path = tmp_path / "missing.md"
    source_path.write_text("![Missing figure](figures/missing figure.png)\n", encoding="utf-8")

    output_path = tmp_path / "missing.docx"

    try:
        markdown_to_docx(source_path, output_path)
    except FileNotFoundError as exc:
        assert "Image not found" in str(exc)
        assert "missing figure.png" in str(exc)
    else:  # pragma: no cover - defensive branch for the assertion style
        raise AssertionError("Expected FileNotFoundError for missing image path")


def test_markdown_to_docx_does_not_duplicate_caption_from_alt_text(tmp_path: Path) -> None:
    source_dir = tmp_path / "notes"
    image_dir = source_dir / "figures"
    source_dir.mkdir()
    image_dir.mkdir()

    image_path = image_dir / "stage 1.png"
    image_path.write_bytes(_ONE_PIXEL_PNG)

    source_path = source_dir / "report.md"
    source_path.write_text(
        "# Thesis Report\n\n![Figure 1 caption.](figures/stage 1.png)\n\nFigure 1 caption.\n",
        encoding="utf-8",
    )

    output_path = tmp_path / "report.docx"
    markdown_to_docx(source_path, output_path)

    document = Document(output_path)
    assert _non_empty_paragraph_texts(document) == [
        "Thesis Report",
        "Figure 1 caption.",
    ]
    assert len(document.inline_shapes) == 1


def test_markdown_to_docx_uses_quoted_image_title_as_caption(tmp_path: Path) -> None:
    source_dir = tmp_path / "notes"
    image_dir = source_dir / "figures"
    source_dir.mkdir()
    image_dir.mkdir()

    image_path = image_dir / "stage 1.png"
    image_path.write_bytes(_ONE_PIXEL_PNG)

    source_path = source_dir / "report.md"
    source_path.write_text(
        '# Thesis Report\n\n![Alt text](figures/stage 1.png "Figure 1. Caption")\n',
        encoding="utf-8",
    )

    output_path = tmp_path / "report.docx"
    markdown_to_docx(source_path, output_path)

    document = Document(output_path)
    assert _non_empty_paragraph_texts(document) == [
        "Thesis Report",
        "Figure 1. Caption",
    ]
    assert _paragraph_with_text(document, "Figure 1. Caption").alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert len(document.inline_shapes) == 1


def test_markdown_to_docx_supports_angle_bracket_image_paths_with_titles(tmp_path: Path) -> None:
    source_dir = tmp_path / "notes"
    image_dir = source_dir / "figures"
    source_dir.mkdir()
    image_dir.mkdir()

    image_path = image_dir / "stage 1.png"
    image_path.write_bytes(_ONE_PIXEL_PNG)

    source_path = source_dir / "report.md"
    source_path.write_text(
        '# Thesis Report\n\n![Alt text](<figures/stage 1.png> "Figure 1")\n',
        encoding="utf-8",
    )

    output_path = tmp_path / "report.docx"
    markdown_to_docx(source_path, output_path)

    document = Document(output_path)
    assert _non_empty_paragraph_texts(document) == [
        "Thesis Report",
        "Figure 1",
    ]
    assert len(document.inline_shapes) == 1
    assert _paragraph_with_text(document, "Figure 1").alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_generate_report_docx_cli_exports_heading_bullet_image_and_caption() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    with tempfile.TemporaryDirectory(dir=repo_root) as temp_dir:
        temp_root = Path(temp_dir)
        source_dir = temp_root / "notes"
        image_dir = source_dir / "figures"
        source_dir.mkdir(parents=True)
        image_dir.mkdir()

        image_path = image_dir / "stage 1.png"
        image_path.write_bytes(_ONE_PIXEL_PNG)

        source_path = source_dir / "report.md"
        source_path.write_text(
            "# Report Title\n\n## Findings\n\n- First bullet\n\n![Alt text](<figures/stage 1.png> \"Figure 1. Caption\")\n",
            encoding="utf-8",
        )

        output_path = temp_root / "report.docx"
        relative_source_path = source_path.relative_to(repo_root)
        relative_output_path = output_path.relative_to(repo_root)

        env = os.environ.copy()
        env.pop("PYTHONPATH", None)
        result = subprocess.run(
            [
                sys.executable,
                str(Path("scripts") / "generate_report_docx.py"),
                str(relative_source_path),
                str(relative_output_path),
            ],
            cwd=repo_root,
            env=env,
            capture_output=True,
            text=True,
        )

        assert result.returncode == 0, result.stderr

        document = Document(output_path)
        assert output_path.exists()
        assert any(paragraph.style.name == "Title" and paragraph.text == "Report Title" for paragraph in document.paragraphs)
        assert any(paragraph.style.name == "Heading 1" and paragraph.text == "Findings" for paragraph in document.paragraphs)
        assert any(paragraph.style.name == "List Bullet" and paragraph.text == "First bullet" for paragraph in document.paragraphs)
        assert _non_empty_paragraph_texts(document) == [
            "Report Title",
            "Findings",
            "First bullet",
            "Figure 1. Caption",
        ]
        assert _paragraph_with_text(document, "Figure 1. Caption").alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert len(document.inline_shapes) == 1
